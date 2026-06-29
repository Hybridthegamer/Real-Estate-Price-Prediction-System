"""
generate_report.py
Generates FYP_Chapters_4_5_Appendix.docx containing:
  - Chapter Four: System Implementation
  - Chapter Five: Conclusion
  - Appendix: Source Code, Interfaces, Charts
Font: Times New Roman  |  Guideline: B.Sc Computer Science (Dr. O.E. Taylor)
"""

import os, sys, io, json, textwrap
import numpy as np
import pandas as pd
import joblib
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from matplotlib.gridspec import GridSpec
import matplotlib.patches as patches

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS_DIR = os.path.join(BASE_DIR, 'screenshots')
OUTPUT_FILE = os.path.join(BASE_DIR, 'FYP_Chapters_4_5_Appendix.docx')
FONT = 'Times New Roman'

# ─────────────────────────────────────────────────────────────────────
# ACTUAL METRICS FROM TRAINING RUN
# ─────────────────────────────────────────────────────────────────────
METRICS = {
    'Linear\nRegression': {
        'R2': 0.9518, 'MAE': 11_926_475, 'RMSE': 16_640_296,
        'MAPE': 39.30, 'CV_R2': 0.9487, 'CV_std': 0.0076
    },
    'Decision\nTree': {
        'R2': 0.8392, 'MAE': 15_239_989, 'RMSE': 30_398_767,
        'MAPE': 27.06, 'CV_R2': 0.8977, 'CV_std': 0.0264
    },
    'Random\nForest': {
        'R2': 0.8948, 'MAE': 11_326_172, 'RMSE': 24_582_381,
        'MAPE': 17.16, 'CV_R2': 0.9275, 'CV_std': 0.0264
    },
    'XGBoost\n★': {
        'R2': 0.9369, 'MAE': 9_171_487, 'RMSE': 19_039_025,
        'MAPE': 14.05, 'CV_R2': 0.9550, 'CV_std': 0.0128
    },
}

FEATURE_IMPORTANCES = [
    ('Neighbourhood (Lekki Ph1)', 24.35),
    ('Floor Area (sqft_living)', 12.46),
    ('Neighbourhood (Maitama)', 9.50),
    ('No. of Bathrooms', 9.15),
    ('Neighbourhood (Garki)', 9.14),
    ('Neighbourhood (Lekki Ph2)', 7.15),
    ('Log Floor Area (log_sqft)', 6.62),
    ('Neighbourhood (Ikoyi)', 5.61),
    ('No. of Bedrooms', 4.40),
    ('Neighbourhood (Rumuola)', 3.93),
]

# ─────────────────────────────────────────────────────────────────────
# DOCUMENT STYLE HELPERS
# ─────────────────────────────────────────────────────────────────────

def _set_paragraph_format(para, space_before=0, space_after=8,
                           line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.alignment = alignment


def _set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force Times New Roman for all character sets
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:cs'), FONT)
    rPr.insert(0, rFonts)


def set_document_margins(doc):
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.0)


def chapter_heading(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(para, space_before=12, space_after=20, line_spacing=1.0)
    run = para.add_run(text.upper())
    _set_run_font(run, size=14, bold=True)
    return para


def section_heading(doc, text):
    para = doc.add_paragraph()
    _set_paragraph_format(para, space_before=12, space_after=6, line_spacing=1.0,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    _set_run_font(run, size=12, bold=True)
    return para


def subsection_heading(doc, text):
    para = doc.add_paragraph()
    _set_paragraph_format(para, space_before=8, space_after=4, line_spacing=1.0,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    _set_run_font(run, size=12, bold=True, italic=True)
    return para


def body(doc, text, space_after=8):
    para = doc.add_paragraph()
    _set_paragraph_format(para, space_before=0, space_after=space_after)
    para.paragraph_format.first_line_indent = Inches(0.5)
    run = para.add_run(text)
    _set_run_font(run, size=12)
    return para


def body_no_indent(doc, text, space_after=8):
    para = doc.add_paragraph()
    _set_paragraph_format(para, space_before=0, space_after=space_after)
    run = para.add_run(text)
    _set_run_font(run, size=12)
    return para


def bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    _set_paragraph_format(para, space_before=0, space_after=4, line_spacing=1.5,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT)
    para.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    run = para.add_run(text)
    _set_run_font(run, size=12)
    return para


def numbered(doc, text, level=0):
    para = doc.add_paragraph(style='List Number')
    _set_paragraph_format(para, space_before=0, space_after=4, line_spacing=1.5,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT)
    para.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    run = para.add_run(text)
    _set_run_font(run, size=12)
    return para


def figure_caption(doc, text):
    para = doc.add_paragraph()
    _set_paragraph_format(para, space_before=4, space_after=12, line_spacing=1.0,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = para.add_run(text)
    _set_run_font(run, size=11, italic=True, bold=True)
    return para


def table_caption(doc, text):
    para = doc.add_paragraph()
    _set_paragraph_format(para, space_before=12, space_after=4, line_spacing=1.0,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = para.add_run(text)
    _set_run_font(run, size=11, italic=True, bold=True)
    return para


def add_image(doc, img_bytes_or_path, caption, width=5.5):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(para, space_before=6, space_after=2, line_spacing=1.0)
    if isinstance(img_bytes_or_path, bytes):
        run = para.add_run()
        run.add_picture(io.BytesIO(img_bytes_or_path), width=Inches(width))
    elif isinstance(img_bytes_or_path, io.BytesIO):
        img_bytes_or_path.seek(0)
        run = para.add_run()
        run.add_picture(img_bytes_or_path, width=Inches(width))
    elif isinstance(img_bytes_or_path, str) and os.path.exists(img_bytes_or_path):
        run = para.add_run()
        run.add_picture(img_bytes_or_path, width=Inches(width))
    figure_caption(doc, caption)


def add_table(doc, headers, rows, caption_text=''):
    if caption_text:
        table_caption(doc, caption_text)
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = 'Table Grid'
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        _set_run_font(run, size=11, bold=True)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F3864')
        shading.set(qn('w:color'), 'FFFFFF')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            bold_row = r_idx == len(rows) - 1  # highlight last row (best model)
            run = para.add_run(str(cell_text))
            _set_run_font(run, size=11, bold=bold_row)
    doc.add_paragraph()  # spacing after table


def add_code(doc, code_text, caption=''):
    if caption:
        para = doc.add_paragraph()
        _set_paragraph_format(para, space_before=8, space_after=2, line_spacing=1.0,
                               alignment=WD_ALIGN_PARAGRAPH.LEFT)
        run = para.add_run(caption)
        _set_run_font(run, size=11, bold=True, italic=True)

    # Shade the code block
    para = doc.add_paragraph()
    _set_paragraph_format(para, space_before=0, space_after=0, line_spacing=1.0,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Grey background for code
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────
# CHART GENERATORS
# ─────────────────────────────────────────────────────────────────────

COLORS = ['#5B7FA6', '#7EA6C4', '#3A6EA5', '#1F3864']
BEST_COLOR = '#1F3864'

def _buf(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def chart_model_comparison():
    names = list(METRICS.keys())
    r2    = [METRICS[n]['R2']   for n in names]
    mape  = [METRICS[n]['MAPE'] for n in names]
    mae   = [METRICS[n]['MAE'] / 1_000_000 for n in names]

    fig, axes = plt.subplots(1, 3, figsize=(13, 4.5))
    fig.suptitle('Figure 4.1: Comparative Model Performance on Hold-Out Test Set',
                 fontsize=11, fontweight='bold', y=1.02)

    bar_colors = [COLORS[i] for i in range(len(names))]
    bar_colors[-1] = BEST_COLOR

    for ax, vals, ylabel, title, fmt in zip(
        axes,
        [r2, mape, mae],
        ['R² Score', 'MAPE (%)', 'MAE (₦ Millions)'],
        ['R² (higher is better)', 'MAPE – % Error (lower is better)',
         'MAE (lower is better)'],
        ['.4f', '.1f', '.1f']
    ):
        bars = ax.bar(names, vals, color=bar_colors, edgecolor='white', linewidth=0.5,
                      zorder=3)
        ax.set_title(title, fontsize=9, fontweight='bold', pad=6)
        ax.set_ylabel(ylabel, fontsize=8)
        ax.tick_params(axis='x', labelsize=7.5)
        ax.tick_params(axis='y', labelsize=8)
        ax.grid(axis='y', alpha=0.3, zorder=0)
        ax.set_axisbelow(True)
        for bar, val in zip(bars, vals):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(vals)*0.01,
                    f'{val:{fmt}}', ha='center', va='bottom', fontsize=8, fontweight='bold')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    plt.tight_layout()
    return _buf(fig)


def chart_cv_r2():
    names  = list(METRICS.keys())
    means  = [METRICS[n]['CV_R2'] for n in names]
    stds   = [METRICS[n]['CV_std'] for n in names]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    bar_colors = [COLORS[i] for i in range(len(names))]
    bar_colors[-1] = BEST_COLOR

    bars = ax.bar(names, means, yerr=stds, capsize=5, color=bar_colors,
                  edgecolor='white', linewidth=0.5, zorder=3,
                  error_kw={'elinewidth': 1.5, 'ecolor': '#555'})
    ax.set_title('Figure 4.2: 10-Fold Cross-Validation R² Scores (Mean ± Std)',
                 fontsize=10, fontweight='bold', pad=10)
    ax.set_ylabel('Cross-Validated R²', fontsize=10)
    ax.set_ylim(0.80, 1.0)
    ax.tick_params(axis='x', labelsize=9)
    ax.grid(axis='y', alpha=0.3, zorder=0)
    ax.set_axisbelow(True)
    for bar, m, s in zip(bars, means, stds):
        ax.text(bar.get_x() + bar.get_width()/2, m + s + 0.002,
                f'{m:.4f}', ha='center', va='bottom', fontsize=9, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.annotate('★ Selected', xy=(bars[-1].get_x() + bars[-1].get_width()/2, means[-1]),
                xytext=(3.2, 0.97), fontsize=9, color=BEST_COLOR, fontweight='bold',
                arrowprops=dict(arrowstyle='->', color=BEST_COLOR))
    plt.tight_layout()
    return _buf(fig)


def chart_feature_importance():
    labels = [f[0] for f in FEATURE_IMPORTANCES]
    values = [f[1] for f in FEATURE_IMPORTANCES]
    colors = [BEST_COLOR if i < 3 else '#5B7FA6' for i in range(len(labels))]

    fig, ax = plt.subplots(figsize=(9, 5))
    bars = ax.barh(labels[::-1], values[::-1], color=colors[::-1],
                   edgecolor='white', linewidth=0.4)
    ax.set_xlabel('Relative Importance (%)', fontsize=10)
    ax.set_title('Figure 4.3: Top-10 XGBoost Feature Importances',
                 fontsize=10, fontweight='bold', pad=10)
    ax.tick_params(axis='y', labelsize=9)
    ax.tick_params(axis='x', labelsize=9)
    ax.grid(axis='x', alpha=0.3)
    ax.set_axisbelow(True)
    for bar, val in zip(bars, values[::-1]):
        ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2,
                f'{val:.2f}%', va='center', fontsize=8.5, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    return _buf(fig)


def chart_price_distribution():
    df = pd.read_csv(os.path.join(BASE_DIR, 'data', 'raw', 'housing_data.csv'))
    prices_m = df['price'] / 1_000_000

    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))
    fig.suptitle('Figure 4.4: Training Dataset Price Distribution',
                 fontsize=10, fontweight='bold')

    axes[0].hist(prices_m, bins=60, color=BEST_COLOR, edgecolor='white',
                 linewidth=0.3, alpha=0.85)
    axes[0].set_xlabel('Price (₦ Millions)', fontsize=9)
    axes[0].set_ylabel('Frequency', fontsize=9)
    axes[0].set_title('Full Price Distribution', fontsize=9, fontweight='bold')
    axes[0].set_xlim(0, 600)
    axes[0].grid(axis='y', alpha=0.3)
    axes[0].spines['top'].set_visible(False)
    axes[0].spines['right'].set_visible(False)

    # Box plot by property type
    types = df['property_type'].unique()
    data_by_type = [df[df['property_type'] == t]['price'].values / 1_000_000 for t in types]
    bp = axes[1].boxplot(data_by_type, patch_artist=True, showfliers=False,
                          medianprops={'color': 'white', 'linewidth': 2})
    for patch, color in zip(bp['boxes'], [COLORS[i % len(COLORS)] for i in range(len(types))]):
        patch.set_facecolor(color)
    axes[1].set_xticklabels([t.replace(' ', '\n') for t in types], fontsize=7.5)
    axes[1].set_ylabel('Price (₦ Millions)', fontsize=9)
    axes[1].set_title('Price by Property Type (Outliers Hidden)', fontsize=9, fontweight='bold')
    axes[1].grid(axis='y', alpha=0.3)
    axes[1].spines['top'].set_visible(False)
    axes[1].spines['right'].set_visible(False)

    plt.tight_layout()
    return _buf(fig)


def chart_predicted_vs_actual():
    pipeline = joblib.load(os.path.join(BASE_DIR, 'models', 'model_pipeline.pkl'))
    df_proc = pd.read_csv(os.path.join(BASE_DIR, 'data', 'processed', 'processed_data.csv'))
    sys.path.insert(0, BASE_DIR)
    from src.feature_engineering import engineer_features
    df_eng = engineer_features(df_proc)
    from sklearn.model_selection import train_test_split
    X = df_eng.drop(columns=['price'])
    y = df_eng['price']
    _, X_test, _, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
    y_pred = pipeline.predict(X_test)

    y_test_m = np.array(y_test) / 1_000_000
    y_pred_m = y_pred / 1_000_000

    fig, ax = plt.subplots(figsize=(7, 6))
    ax.scatter(y_test_m, y_pred_m, alpha=0.40, s=18, color=BEST_COLOR, edgecolors='none')
    lim = max(y_test_m.max(), y_pred_m.max()) * 1.05
    ax.plot([0, lim], [0, lim], 'r--', linewidth=1.5, label='Perfect Prediction (y = x)')
    ax.set_xlabel('Actual Price (₦ Millions)', fontsize=10)
    ax.set_ylabel('Predicted Price (₦ Millions)', fontsize=10)
    ax.set_title('Figure 4.5: XGBoost – Predicted vs. Actual Prices\n(Hold-Out Test Set, n=533)',
                 fontsize=10, fontweight='bold')
    ax.legend(fontsize=9)
    ax.grid(alpha=0.25)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_xlim(0, lim)
    ax.set_ylim(0, lim)
    # R² annotation
    from sklearn.metrics import r2_score
    r2 = r2_score(y_test, y_pred)
    ax.text(0.05, 0.93, f'R² = {r2:.4f}', transform=ax.transAxes,
            fontsize=10, fontweight='bold', color=BEST_COLOR,
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#E8EFF8', alpha=0.8))
    plt.tight_layout()
    return _buf(fig)


# ─────────────────────────────────────────────────────────────────────
# SOURCE CODE READER
# ─────────────────────────────────────────────────────────────────────

def read_source(filename):
    path = os.path.join(BASE_DIR, filename)
    if not os.path.exists(path):
        return f'# File not found: {filename}\n'
    with open(path, 'r') as f:
        return f.read()


# ─────────────────────────────────────────────────────────────────────
# CHAPTER FOUR
# ─────────────────────────────────────────────────────────────────────

def build_chapter_four(doc):

    chapter_heading(doc, 'Chapter Four')
    chapter_heading(doc, 'System Implementation')

    # ── 4.1 System Implementation ──────────────────────────────────
    section_heading(doc, '4.1 System Implementation and Implementation Results')

    body(doc,
        'The Real Estate Price Prediction System was implemented in full accordance '
        'with the design specifications detailed in Chapter Three, adopting the Rapid '
        'Application Development (RAD) methodology as the governing software development '
        'framework. The RAD approach enabled iterative, feedback-driven refinement of both '
        'the machine learning pipeline and the web application interface across five '
        'sequential development phases: (1) environment configuration and project '
        'scaffolding; (2) data pipeline development; (3) machine learning model '
        'development and evaluation; (4) web application and API development; and '
        '(5) integration, testing, and model persistence. This chapter presents a detailed '
        'account of each implementation phase, the outcomes achieved, the system outputs '
        'obtained through testing, and the rationale underlying the selection of '
        'development tools and platforms.')

    body(doc,
        'The system is entirely implemented in the Python programming language (version '
        '3.11), leveraging a carefully curated suite of open-source scientific computing '
        'and web development libraries. The Flask micro-framework (version 3.1.3) serves '
        'as the web application backbone, while scikit-learn (version 1.9.0) and XGBoost '
        '(version 3.2.0) provide the machine learning modelling capabilities. Data '
        'manipulation and numerical computation are handled by Pandas (version 3.0.3) and '
        'NumPy (version 2.4.6) respectively, and model artefact persistence is managed '
        'using Joblib (version 1.5.3). The front-end interface is constructed using '
        'Bootstrap 5.3.2 for responsive layout and Chart.js 4.4.0 for interactive data '
        'visualisation. The complete source code is modularised into independent, '
        'loosely-coupled packages, with each module addressing a distinct computational '
        'concern (preprocessing, feature engineering, training, inference), in adherence '
        'to the PEP 8 coding standards mandated by the system\'s non-functional '
        'requirements.')

    # ── 4.2 Data Pipeline ──────────────────────────────────────────
    section_heading(doc, '4.2 Data Pipeline Implementation')

    subsection_heading(doc, '4.2.1 Dataset Generation and Collection')
    body(doc,
        'Given the well-documented scarcity of publicly available, structured Nigerian '
        'residential property transaction records (Olanrele et al., 2019), the system\'s '
        'training dataset was generated synthetically using a domain-informed price model '
        'calibrated to reflect realistic price distributions across twenty-six (26) '
        'residential neighbourhoods in three major Nigerian cities — Lagos (twelve '
        'neighbourhoods), Abuja (eight neighbourhoods), and Port Harcourt (six '
        'neighbourhoods). The price model incorporated location tier classifications '
        '(ultra-premium, premium, mid-high, and standard) with corresponding base prices '
        'per square metre ranging from ₦80,000 (Eleme, Port Harcourt) to ₦2,500,000 '
        '(Banana Island, Lagos). A multiplicative pricing function incorporating property '
        'type, age, construction grade, physical condition, and amenity presence was '
        'applied to each record, with Gaussian noise (standard deviation = 8%) added to '
        'simulate realistic market variability.')

    body(doc,
        'A total of 3,000 property records were generated and stored in '
        'data/raw/housing_data.csv. The dataset comprised the following '
        'distribution of key attributes: property types — Detached Houses (22%), '
        'Apartments (21%), Semi-Detached Houses (20%), Terraced Houses (19%), and '
        'Bungalows (18%); floor area ranging from 40 to 1,500 m² (mean: 189 m²); '
        'bedrooms ranging from one to six (mean: 3.2); and sale prices ranging from '
        '₦5,000,000 to ₦1,745,130,000 (median: ₦51,110,000). This broad price range '
        'captures the stark socioeconomic stratification of the Nigerian property market, '
        'from affordable bungalows in peripheral urban areas to luxury detached residences '
        'in premium central districts.')

    subsection_heading(doc, '4.2.2 Data Preprocessing')
    body(doc,
        'The raw dataset was passed through the preprocessing pipeline implemented in '
        'src/data_preprocessing.py, executing the steps specified in Algorithm 3.1 '
        '(Chapter Three). Duplicate record detection yielded no duplicates in the '
        'synthetic dataset. Missing value imputation (median strategy) was confirmed '
        'functional but found no missing entries. Outlier removal using the Inter-Quartile '
        'Range (IQR) method with a threshold of 1.5 identified 335 records (11.2% of the '
        'dataset) as statistical outliers across the target price and floor area '
        'distributions; these were removed to yield a cleaned dataset of 2,665 records '
        'for model training and evaluation. The cleaned dataset was persisted to '
        'data/processed/processed_data.csv for reproducibility.')

    subsection_heading(doc, '4.2.3 Feature Engineering')
    body(doc,
        'Four derived features were engineered from the raw dataset attributes, as '
        'specified in the ENGINEER_FEATURES step of Algorithm 3.1:')

    bullet(doc, 'property_age: Computed as (2025 – yr_built), converting the raw '
           'construction year to an interpretable age metric. Range: 0 – 45 years.')
    bullet(doc, 'renovation_flag: Binary indicator (1 if yr_renovated > 0, otherwise 0). '
           'Approximately 25.3% of records in the cleaned dataset indicated a renovation '
           'history.')
    bullet(doc, 'log_sqft: The natural logarithm of sqft_living (floor area), applied to '
           'reduce the right-skewness of the floor area distribution and approximately '
           'linearise the price-to-area relationship in the feature space.')
    bullet(doc, 'bedroom_bathroom_ratio: The ratio of bedrooms to bathrooms, capturing the '
           'internal space efficiency and composition of each property.')

    body(doc,
        'The raw temporal columns yr_built and yr_renovated were dropped after feature '
        'engineering to prevent data redundancy. The final processed and engineered '
        'dataset comprised 2,665 records and twenty (20) feature columns prior to '
        'categorical encoding.')

    # ── 4.3 Model Training ──────────────────────────────────────────
    section_heading(doc, '4.3 Machine Learning Model Implementation and Training')

    body(doc,
        'The machine learning training pipeline was implemented in '
        'src/model_training.py using scikit-learn\'s Pipeline and ColumnTransformer '
        'APIs (Pedregosa et al., 2011). This design bundles the preprocessing '
        'transformations — StandardScaler for numerical features and OneHotEncoder '
        '(with handle_unknown=\'ignore\') for categorical features — together with the '
        'regression model into a single, serialisable Python object. This architectural '
        'decision is critical for inference correctness, as it guarantees that the '
        'identical transformations fitted on the training data are automatically reapplied '
        'to all runtime prediction inputs, eliminating the risk of data leakage or '
        'transform inconsistency between training and serving environments.')

    body(doc,
        'The 2,665 cleaned records were partitioned into a training set (80%, 2,132 '
        'records) and a hold-out test set (20%, 533 records) using a fixed random seed '
        '(random_state=42) for reproducibility. Four regression algorithms were trained '
        'and evaluated:')

    bullet(doc, 'Linear Regression (baseline model, Ordinary Least Squares via scikit-learn '
           'LinearRegression())')
    bullet(doc, 'Decision Tree Regressor (max_depth=10, min_samples_split=10)')
    bullet(doc, 'Random Forest Regressor (n_estimators=200, max_depth=15, '
           'min_samples_split=5; Breiman, 2001)')
    bullet(doc, 'XGBoost Regressor (n_estimators=200, learning_rate=0.1, max_depth=6, '
           'L1 regularisation α=0.1, L2 regularisation λ=1.0; Chen & Guestrin, 2016)')

    body(doc,
        'Each fitted pipeline was additionally evaluated using 10-fold cross-validation '
        'on the training set (Algorithm 3.3, Chapter Three), reporting mean CV R² and its '
        'standard deviation across the ten folds. Model selection was based on '
        'cross-validated R², the statistically principled criterion that mitigates '
        'test-set overfitting and provides the most reliable estimate of generalisation '
        'performance.')

    # ── 4.4 Implementation Results ──────────────────────────────────
    section_heading(doc, '4.4 Implementation Results')

    subsection_heading(doc, '4.4.1 Model Performance Comparison')
    body(doc,
        'Table 4.1 presents the performance metrics for all four regression models '
        'evaluated on the 20% hold-out test set, together with 10-fold cross-validation '
        'results. Figure 4.1 provides a comparative visualisation of the three primary '
        'performance indicators.')

    table_caption(doc, 'Table 4.1: Model Performance Comparison on Hold-Out Test Set '
                       '(n = 533) and 10-Fold Cross-Validation')
    add_table(doc,
        headers=['Algorithm', 'Test R²', 'Test MAE (₦)', 'Test RMSE (₦)', 'MAPE (%)', 'CV R² (Mean ± Std)'],
        rows=[
            ['Linear Regression', '0.9518', '11,926,475', '16,640,296', '39.30', '0.9487 ± 0.0076'],
            ['Decision Tree',     '0.8392', '15,239,989', '30,398,767', '27.06', '0.8977 ± 0.0264'],
            ['Random Forest',     '0.8948', '11,326,172', '24,582,381', '17.16', '0.9275 ± 0.0264'],
            ['XGBoost ★',         '0.9369',  '9,171,487', '19,039,025', '14.05', '0.9550 ± 0.0128'],
        ]
    )

    add_image(doc, chart_model_comparison(),
              'Figure 4.1: Comparative Model Performance – R², MAPE, and MAE (Test Set)',
              width=5.8)

    body(doc,
        'The results in Table 4.1 reveal several significant findings. While Linear '
        'Regression achieved the highest test-set R² (0.9518), it exhibited a '
        'Mean Absolute Percentage Error (MAPE) of 39.30%, indicating severe proportional '
        'prediction errors — particularly for lower-priced properties where the linear '
        'assumption between features and price is substantially violated. This disparity '
        'between R² and MAPE in the linear model highlights a known limitation of R² as a '
        'sole evaluation criterion: a model may explain a large proportion of variance '
        'while still producing practically unacceptable percentage errors (Willmott & '
        'Matsuura, 2005).')

    body(doc,
        'XGBoost delivered the most practically useful performance profile across all '
        'evaluation criteria of primary practical importance. Its MAPE of 14.05% '
        'represents the lowest proportional prediction error of all four models — nearly '
        'three times lower than Linear Regression and 18% lower than Random Forest — '
        'indicating that the median XGBoost prediction is within approximately 14% of the '
        'actual market price. Crucially, XGBoost also achieved the highest cross-validated '
        'R² (0.9550 ± 0.0128), confirming that its generalisation performance is robust '
        'and consistent across different subsets of the training data. The considerably '
        'lower standard deviation of XGBoost\'s CV R² scores (0.0128) compared to '
        'Random Forest (0.0264) and Decision Tree (0.0264) further demonstrates that '
        'XGBoost produces more stable predictions across varying data splits, a direct '
        'consequence of its L1 and L2 regularisation terms controlling model complexity '
        '(Chen & Guestrin, 2016).')

    add_image(doc, chart_cv_r2(),
              'Figure 4.2: 10-Fold Cross-Validation R² Scores with Standard Deviation Error Bars',
              width=5.0)

    subsection_heading(doc, '4.4.2 Model Selection and Deployment')
    body(doc,
        'Based on the evaluation results, XGBoost was selected as the deployed production '
        'model. The selection criterion — cross-validated R² — is the statistically '
        'principled choice for model selection, as it estimates generalisation performance '
        'without relying on the test set (which would constitute data leakage into the '
        'selection decision). The trained XGBoost pipeline, encapsulating the fitted '
        'ColumnTransformer preprocessor and XGBoost regressor, was serialised to the '
        'binary file models/model_pipeline.pkl using Joblib (version 1.5.3). This '
        'serialised pipeline object is loaded into memory once at application startup '
        'and reused for all subsequent prediction requests, yielding sub-second inference '
        'latency consistent with the system\'s 2-second response time requirement.')

    body(doc,
        'The model comparison metadata, including all performance metrics and the '
        'selection rationale, was persisted to models/metrics.json and is displayed '
        'in the Administration Panel, enabling authorised administrators to review '
        'model performance at any time without re-running the training pipeline. '
        'A complete model retraining cycle — including data preprocessing, feature '
        'engineering, training all four algorithms, cross-validation, model selection, '
        'and serialisation — completes in approximately 2 to 5 minutes on a standard '
        'laptop CPU, satisfying the maintainability requirement of the system.')

    subsection_heading(doc, '4.4.3 Feature Importance Analysis')
    body(doc,
        'A key advantage of ensemble tree-based models such as XGBoost over linear '
        'regression is their capacity to directly provide feature importance estimates, '
        'quantifying the relative contribution of each input attribute to the overall '
        'predictive performance of the model. Figure 4.3 presents the top ten most '
        'influential features as extracted from the deployed XGBoost model via its '
        'feature_importances_ attribute.')

    add_image(doc, chart_feature_importance(),
              'Figure 4.3: Top-10 Feature Importances of Deployed XGBoost Model',
              width=5.5)

    body(doc,
        'The feature importance analysis reveals that location — represented by the '
        'one-hot encoded neighbourhood indicators — is by far the dominant predictor '
        'of residential property price, consistent with the hedonic pricing theory '
        'reviewed in Chapter Two (Rosen, 1974; Mu et al., 2022). The Lekki Phase 1 '
        'neighbourhood indicator alone accounts for 24.35% of the model\'s predictive '
        'weight, reflecting the premium commanded by properties in this high-demand '
        'Lagos coastal neighbourhood. Floor area (sqft_living, 12.46%) is the second '
        'most influential feature, confirming the intuitive expectation that larger '
        'properties attract higher prices. The Maitama, Garki, and Lekki Phase 2 '
        'neighbourhood indicators collectively contribute a further 25% of the model\'s '
        'predictive weight, reinforcing the primacy of location in property valuation. '
        'The number of bathrooms (9.15%) ranks ahead of bedrooms (4.40%), suggesting '
        'that bathroom provision is a stronger signal of property quality and market '
        'tier than bedroom count alone in the Nigerian residential market context.')

    add_image(doc, chart_price_distribution(),
              'Figure 4.4: Training Dataset Price Distribution and Price by Property Type',
              width=5.8)

    add_image(doc, chart_predicted_vs_actual(),
              'Figure 4.5: XGBoost Predicted vs. Actual Prices on Hold-Out Test Set (n = 533)',
              width=4.5)

    body(doc,
        'Figure 4.5 presents the scatter plot of XGBoost predicted prices versus actual '
        'prices on the 533-record hold-out test set. The concentration of data points '
        'along the perfect prediction diagonal (y = x) demonstrates that the model '
        'produces broadly accurate predictions across the full price range. Modest '
        'dispersion is visible at higher price points (above ₦200 million), where the '
        'reduced frequency of ultra-premium training examples limits the model\'s '
        'extrapolation capacity — a characteristic expected of data-driven models '
        'trained on datasets with inherent class imbalance.')

    # ── 4.5 Sample Outputs ──────────────────────────────────────────
    section_heading(doc, '4.5 Sample Outputs')

    body(doc,
        'This section presents screenshots of the four principal user interfaces of the '
        'Real Estate Price Prediction System as rendered during functional testing. The '
        'application was accessed via Google Chrome in a local development environment '
        'at http://localhost:5000. All interface elements conform to the input/output '
        'design specifications detailed in Chapter Three (§3.5).')

    subsection_heading(doc, '4.5.1 Home / Landing Page')
    body(doc,
        'The home page (Figure 4.6) serves as the system\'s primary entry point. It '
        'presents a hero section with a clear value proposition, navigation links to the '
        'prediction form, an indicative price summary card for key Nigerian cities, '
        'feature highlight cards, a three-step "How It Works" guide, and a model '
        'performance comparison table. The interface is rendered using the Bootstrap 5 '
        'responsive grid, with the deep blue (#1F3864) primary colour scheme and amber '
        'accent (#F5A623) specified in Chapter Three, §3.12.')

    if os.path.exists(os.path.join(SCREENSHOTS_DIR, 'home.png')):
        add_image(doc, os.path.join(SCREENSHOTS_DIR, 'home.png'),
                  'Figure 4.6: EstimateNG Landing Page (Home Screen)',
                  width=5.5)

    subsection_heading(doc, '4.5.2 Property Prediction Form')
    body(doc,
        'The prediction form page (Figure 4.7) presents the property attribute input '
        'interface structured into three logically grouped, clearly labelled sections: '
        'Location Details (Neighbourhood/City and Property Type), Property Specifications '
        '(Floor Area, Year of Construction, Bedrooms, Bathrooms, and Car Parking Spaces), '
        'and Additional Features & Amenities (Swimming Pool, Gated Estate, and '
        'Gym/Fitness Facility checkboxes). All fields implement real-time client-side '
        'validation with inline error messages, consistent with the functional requirement '
        'that "the system shall validate all user inputs for completeness, data type '
        'correctness, and permissible range before submission" (§3.3.2). A submit button '
        'triggers the AJAX prediction request, and a reset button clears all fields '
        'without reloading the page.')

    if os.path.exists(os.path.join(SCREENSHOTS_DIR, 'predict.png')):
        add_image(doc, os.path.join(SCREENSHOTS_DIR, 'predict.png'),
                  'Figure 4.7: Property Prediction Input Form',
                  width=5.5)

    subsection_heading(doc, '4.5.3 Prediction Results Interface')
    body(doc,
        'Upon successful form submission, the prediction results are dynamically rendered '
        'below the input form (Figure 4.8) without a page reload, via an asynchronous '
        'AJAX call to the POST /api/predict endpoint. The results interface displays '
        'three categories of information as specified in §3.5.2: (1) the predicted price '
        'in Nigerian Naira (₦), presented prominently in a styled result card with the '
        'primary colour gradient; (2) the 90% confidence interval as a formatted price '
        'range (e.g., ₦296,120,000 – ₦400,640,000) displayed below the central '
        'estimate; and (3) a feature importance bar chart rendered by Chart.js, showing '
        'the top ten input attributes and their relative percentage contribution to the '
        'prediction. A sample prediction for a 4-bedroom, 3-bathroom, 300 m² Detached '
        'House in Lekki Phase 1 with a swimming pool and gated security, built in 2018, '
        'returned a predicted price of ₦348,380,032, with a 90% confidence range of '
        '₦296,120,000 to ₦400,640,000.')

    subsection_heading(doc, '4.5.4 Administration Panel')
    body(doc,
        'The administration panel (Figure 4.9) is accessible via the protected '
        '/admin route, secured by a session-based password authentication mechanism. '
        'The panel is divided into two primary sections: a left-column panel displaying '
        'the deployed model summary (model name, training date, dataset size, and full '
        'model comparison table) alongside action controls for uploading a new training '
        'dataset and triggering a model retraining cycle; and a right-column panel '
        'presenting a tabular log of the twenty most recent prediction requests, '
        'including timestamp, neighbourhood, property type, floor area, bedroom count, '
        'and predicted price. The prediction log enables ongoing performance monitoring '
        'as required by the functional requirement that "the system shall log each '
        'prediction request...to facilitate model performance monitoring over time" '
        '(§3.3.2). All logged entries are anonymised: no personally identifying '
        'information is retained.')

    if os.path.exists(os.path.join(SCREENSHOTS_DIR, 'admin_login.png')):
        add_image(doc, os.path.join(SCREENSHOTS_DIR, 'admin_login.png'),
                  'Figure 4.9: Administration Panel – Login Screen',
                  width=4.0)

    # ── 4.6 System Setup ──────────────────────────────────────────
    section_heading(doc, '4.6 System Setup (How to Run the Software)')

    subsection_heading(doc, '4.6.1 System Requirements')
    body(doc,
        'The following hardware and software prerequisites are required to deploy and '
        'operate the Real Estate Price Prediction System in a local development '
        'environment:')

    table_caption(doc, 'Table 4.2: Minimum System Requirements')
    add_table(doc,
        headers=['Component', 'Minimum Version / Specification'],
        rows=[
            ['Python', '3.9 or higher'],
            ['pip (Python package manager)', '23.0 or higher'],
            ['RAM', '4 GB (8 GB recommended for model training)'],
            ['Available Disk Space', '500 MB'],
            ['Operating System', 'Windows 10 / macOS 11+ / Ubuntu 20.04+'],
            ['Web Browser', 'Chrome 100+ / Firefox 100+ / Edge 100+ / Safari 15+'],
        ]
    )

    subsection_heading(doc, '4.6.2 Installation and Execution Steps')
    body_no_indent(doc, 'The following steps describe the complete setup and execution procedure:')
    numbered(doc, 'Clone the repository: git clone https://github.com/hybridthegamer/'
             'real-estate-price-prediction-system.git')
    numbered(doc, 'Navigate into the project directory: cd real-estate-price-prediction-system')
    numbered(doc, 'Create and activate a Python virtual environment:')
    bullet(doc, 'Windows: python -m venv venv  then  venv\\Scripts\\activate', level=1)
    bullet(doc, 'macOS / Linux: python3 -m venv venv  then  source venv/bin/activate', level=1)
    numbered(doc, 'Install all Python dependencies: pip install -r requirements.txt')
    numbered(doc, 'Generate the synthetic training dataset: python generate_data.py')
    numbered(doc, 'Execute the model training pipeline: python train.py')
    numbered(doc, 'Launch the Flask web application: python app.py')
    numbered(doc, 'Open a web browser and navigate to: http://localhost:5000')

    body(doc,
        'The model training step (Step 6) executes the complete Algorithm 3.1 pipeline '
        'and typically completes in 2 to 5 minutes on a standard laptop CPU. Upon '
        'completion, the serialised model pipeline is saved to models/model_pipeline.pkl '
        'and the comparison metrics to models/metrics.json. If model_pipeline.pkl '
        'is absent at application startup, the system automatically initiates the '
        'training pipeline before serving any requests, ensuring zero-configuration '
        'operability for first-time deployments.')

    body(doc,
        'The Administration Panel is accessible at http://localhost:5000/admin. '
        'The default administrator password is admin123, which can be overridden '
        'via the ADMIN_PASSWORD environment variable for production deployments. '
        'A dataset CSV file conforming to the schema defined in Table 3.2 (Chapter Three) '
        'can be uploaded via the Admin Panel to replace the training dataset, followed by '
        'a triggered model retraining cycle to update the deployed pipeline.')

    # ── 4.7 Reasons for Platform/Language Choice ──────────────────
    section_heading(doc, '4.7 Reasons for Choice of Platform / Programming Language')

    body(doc,
        'Python was selected as the exclusive development language for this system on '
        'account of its dominant position in the data science and machine learning '
        'ecosystem. Python\'s extensive library ecosystem — encompassing scikit-learn '
        'for machine learning, Pandas and NumPy for data manipulation, Matplotlib and '
        'Seaborn for data visualisation, Joblib for model serialisation, and Flask for '
        'web application development — enables all functional components of the system '
        'to be implemented within a single, coherent programming environment. This '
        'homogeneity reduces architectural complexity and eliminates the need for '
        'inter-language data serialisation interfaces. Python\'s readable, expressive '
        'syntax, first-class support for scientific computing abstractions, and the '
        'PEP 8 coding standard further facilitate code maintainability and readability, '
        'directly satisfying the system\'s maintainability non-functional requirement '
        '(§3.3.3). Python 3.11, the version employed, provides significant performance '
        'improvements over previous major versions through optimised bytecode evaluation, '
        'reducing training and inference latency.')

    body(doc,
        'Flask was chosen as the web application framework for its lightweight, '
        'modular architecture and its minimal learning curve, which are well-suited '
        'to the system\'s relatively constrained web application requirements. Unlike '
        'heavyweight full-stack frameworks such as Django, Flask does not impose an '
        'Object-Relational Mapper (ORM), a fixed project directory structure, or '
        'bundled administrative interfaces. Since the system\'s data layer consists '
        'of flat-file CSV storage and Joblib-serialised model objects rather than a '
        'relational database, Flask\'s lightweight approach avoids unnecessary framework '
        'overhead. Flask\'s native Jinja2 templating engine and straightforward RESTful '
        'routing enabled clean separation between the machine learning inference logic '
        '(src/) and the HTTP request-handling layer (app.py), consistent with the '
        'three-tier architecture specified in Chapter Three (§3.4.1). The Bootstrap 5 '
        'frontend framework was adopted to fulfil the compatibility and responsiveness '
        'requirements of the system across all modern desktop and mobile browsers '
        'without requiring a dedicated frontend build toolchain. Chart.js was selected '
        'for the feature importance visualisation on account of its zero-dependency, '
        'client-side rendering capability, its responsive HTML5 canvas support, and '
        'its clean, professional default aesthetics suitable for a professional '
        'valuation context.')


# ─────────────────────────────────────────────────────────────────────
# CHAPTER FIVE
# ─────────────────────────────────────────────────────────────────────

def build_chapter_five(doc):

    chapter_heading(doc, 'Chapter Five')
    chapter_heading(doc, 'Conclusion')

    section_heading(doc, '5.1 Conclusion')

    body(doc,
        'This project set out to design and implement a machine learning-based Real '
        'Estate Price Prediction System capable of delivering accurate, objective, and '
        'transparent residential property price estimates for the Nigerian market. The '
        'five objectives defined in Chapter One have been fully achieved, as '
        'demonstrated by the implementation evidence presented in Chapter Four. '
        'A synthetic training dataset of 3,000 Nigerian residential property records '
        'was generated using a domain-informed price model encompassing twenty-six '
        'neighbourhoods across Lagos, Abuja, and Port Harcourt. This dataset was '
        'cleaned, preprocessed, and enriched through feature engineering, yielding '
        '2,665 clean records and twenty input features. Four supervised regression '
        'algorithms — Linear Regression, Decision Tree Regression, Random Forest '
        'Regression, and XGBoost — were trained and rigorously evaluated using '
        'hold-out test set metrics (R², MAE, RMSE, MAPE) and 10-fold cross-validation, '
        'precisely as specified by Algorithms 3.1 and 3.3 in Chapter Three.')

    body(doc,
        'XGBoost emerged as the optimal model with a cross-validated R² of 0.9550 '
        '(± 0.0128) and a Mean Absolute Percentage Error of 14.05% on the hold-out '
        'test set, demonstrating that the system\'s predictions are, on average, within '
        '14% of actual market prices — a practically significant improvement over the '
        'manual valuation errors documented in the literature (Olanrele et al., 2019). '
        'The feature importance analysis confirmed the theoretical grounding of the '
        'system in hedonic pricing theory (Rosen, 1974), with neighbourhood location '
        'indicators accounting for over 60% of the model\'s predictive weight, followed '
        'by floor area and bathroom provision as the most influential structural '
        'attributes.')

    body(doc,
        'The trained XGBoost pipeline was deployed as a Flask-powered web application '
        'with a Bootstrap 5 responsive interface, a RESTful prediction API, dynamic '
        'Chart.js feature importance visualisations, session-based administrator '
        'authentication, and an audit log of prediction events. The system delivers '
        'predictions — including a predicted price, a 90% confidence interval, and '
        'a feature importance explanation — in under one second under typical server '
        'load conditions, surpassing the 2-second response time requirement. The '
        'web interface was verified to render correctly on modern desktop browsers '
        'and adapts responsively to mobile screen sizes through Bootstrap\'s '
        'grid system.')

    body(doc,
        'The system directly addresses the three problems identified in Chapter One: '
        'it eliminates the subjectivity of manual appraisals through objective '
        'algorithmic predictions; it removes financial barriers by providing a freely '
        'accessible alternative to costly professional valuations; and it resolves '
        'the inapplicability of foreign AVM platforms to the Nigerian market context '
        'by training on locally-calibrated Nigerian property data. In doing so, '
        'this project makes a concrete contribution to the body of research on '
        'machine learning applications in emerging real estate markets — a gap '
        'explicitly identified in the literature review (Chapter Two, §2.8) — and '
        'provides a replicable, open-source implementation framework that practitioners '
        'and future researchers can adapt and extend.')

    section_heading(doc, '5.2 Recommendations')

    body(doc,
        'While the system achieves its stated objectives with satisfactory predictive '
        'accuracy, several areas of improvement and extension are identified for future '
        'research and development.')

    body(doc,
        'The most impactful improvement would be the replacement or supplementation '
        'of the synthetic training dataset with real Nigerian property transaction '
        'records collected through partnerships with licensed estate surveyors, property '
        'listing platforms (PropertyPro.ng, PrivateProperty.com.ng), or government '
        'land registries. Real transaction data would introduce authentic market noise '
        'patterns, temporal price trends, and socioeconomic dynamics that cannot be '
        'fully simulated, potentially improving both model accuracy and practical '
        'applicability. A data standardisation and collection framework aligned with '
        'the recommendations of Olanrele et al. (2019) is strongly advised as a '
        'precondition for this transition.')

    body(doc,
        'The feature set should be extended to incorporate geospatial proximity '
        'features derived from Geographic Information Systems (GIS) data — '
        'specifically, distances from each property to the nearest primary school, '
        'hospital, public transport hub, shopping centre, and the central business '
        'district. Research by Mu et al. (2022) demonstrated that such geospatial '
        'features improved R² by an average of 13 percentage points across multiple '
        'model architectures in a comparable urban context. Integration of these '
        'features via the Google Maps API or OpenStreetMap Nominatim would be '
        'straightforward given the system\'s existing neighbourhood coordinate '
        'infrastructure.')

    body(doc,
        'The prediction model could be extended to support temporal price forecasting '
        'by incorporating historical sale price time series data. Recurrent Neural '
        'Network architectures — specifically Long Short-Term Memory (LSTM) networks '
        '— have demonstrated strong performance in sequential price forecasting tasks '
        '(Goodfellow et al., 2016) and would complement the current static regression '
        'model by enabling month-over-month or quarter-over-quarter price trajectory '
        'predictions. This would significantly increase the tool\'s utility for '
        'real estate investors and mortgage lenders.')

    body(doc,
        'Model interpretability could be enhanced through the integration of SHAP '
        '(SHapley Additive exPlanations) values (Ribeiro et al., 2016), which provide '
        'individual, prediction-level explanations rather than global feature importance '
        'averages. SHAP values would enable the system to explain, for each specific '
        'prediction, precisely how much each input feature increased or decreased the '
        'predicted price relative to the average prediction — substantially improving '
        'the transparency and trustworthiness of the system for property buyers, '
        'sellers, and lenders.')

    body(doc,
        'From a deployment perspective, the system should be migrated from a single-server '
        'development configuration to a production-grade deployment stack. This would '
        'entail replacing Flask\'s development server with a production WSGI server '
        '(such as Gunicorn), deploying behind an Nginx reverse proxy, containerising '
        'the application using Docker, and hosting on a cloud platform such as AWS, '
        'Google Cloud, or Azure. A PostgreSQL database should replace the current '
        'flat-file storage for the prediction log, enabling structured querying and '
        'long-term performance monitoring. These changes would bring the system into '
        'full alignment with the scalability and availability requirements specified '
        'in §3.3.3.')

    body(doc,
        'Finally, the scope of the system should be expanded beyond residential '
        'property to encompass commercial real estate (office spaces, retail units, '
        'warehouses), which represents a significant and underserved segment of the '
        'Nigerian property market. Separate model pipelines could be trained for '
        'each property category, with shared preprocessing infrastructure, enabling '
        'a comprehensive automated valuation platform covering the full spectrum of '
        'Nigerian real estate asset classes.')


# ─────────────────────────────────────────────────────────────────────
# APPENDIX
# ─────────────────────────────────────────────────────────────────────

def build_appendix(doc):

    chapter_heading(doc, 'Appendices')

    # ── Appendix A: Source Code ────────────────────────────────────
    section_heading(doc, 'Appendix A: System Source Code')
    body_no_indent(doc,
        'The following listings present the complete source code of the primary system '
        'modules. All code is written in Python 3.11 in adherence to PEP 8 coding '
        'standards. The full project repository, including all templates, static files, '
        'and configuration, is available on GitHub.')

    code_files = [
        ('config.py',                    'Listing A.1: config.py – System Configuration and Constants'),
        ('train.py',                     'Listing A.2: train.py – Training Entry Point'),
        ('src/data_preprocessing.py',    'Listing A.3: src/data_preprocessing.py – Data Preprocessing Module'),
        ('src/feature_engineering.py',   'Listing A.4: src/feature_engineering.py – Feature Engineering Module'),
        ('src/model_training.py',        'Listing A.5: src/model_training.py – Model Training and Evaluation'),
        ('src/prediction_engine.py',     'Listing A.6: src/prediction_engine.py – Inference Pipeline'),
        ('app.py',                       'Listing A.7: app.py – Flask Web Application'),
    ]

    for filepath, caption in code_files:
        code = read_source(filepath)
        add_code(doc, code, caption)

    doc.add_page_break()

    # ── Appendix B: JavaScript Frontend ───────────────────────────
    section_heading(doc, 'Appendix B: Frontend Source Code')
    add_code(doc, read_source('static/js/main.js'),
             'Listing B.1: static/js/main.js – Prediction Form JavaScript Handler')

    doc.add_page_break()

    # ── Appendix C: Dataset Schema ─────────────────────────────────
    section_heading(doc, 'Appendix C: Training Dataset Schema')
    body_no_indent(doc,
        'Table C.1 presents the complete schema of the training dataset stored in '
        'data/raw/housing_data.csv, as generated by generate_data.py and defined '
        'in Table 3.2 of Chapter Three.')

    table_caption(doc, 'Table C.1: Full Training Dataset Schema')
    add_table(doc,
        headers=['Column Name', 'Data Type', 'Description', 'Example Value'],
        rows=[
            ['price',                  'FLOAT',   'Target: property sale price (₦)',          '85,500,000'],
            ['neighbourhood',          'STRING',  'Neighbourhood or city area name',          'Lekki Phase 1'],
            ['property_type',          'STRING',  'Residential property category',            'Detached House'],
            ['sqft_living',            'INTEGER', 'Interior living area (m²)',                '320'],
            ['sqft_lot',               'INTEGER', 'Total lot/compound size (m²)',             '650'],
            ['bedrooms',               'INTEGER', 'Number of bedrooms',                       '4'],
            ['bathrooms',              'INTEGER', 'Number of bathrooms',                      '3'],
            ['floors',                 'INTEGER', 'Number of floors in building',             '2'],
            ['parking_spaces',         'INTEGER', 'Number of car parking spaces',             '2'],
            ['yr_built',               'INTEGER', 'Year property was originally built',       '2015'],
            ['yr_renovated',           'INTEGER', 'Year of most recent renovation (0=none)',  '0'],
            ['has_pool',               'INTEGER', 'Binary: 1=has swimming pool',              '1'],
            ['is_gated',               'INTEGER', 'Binary: 1=in gated estate',                '1'],
            ['has_gym',                'INTEGER', 'Binary: 1=has gym/fitness facility',       '0'],
            ['grade',                  'INTEGER', 'Construction quality grade (1–13)',        '8'],
            ['condition',              'INTEGER', 'Overall property condition (1–5)',         '4'],
            ['lat',                    'FLOAT',   'Latitude coordinate',                      '6.4352'],
            ['long',                   'FLOAT',   'Longitude coordinate',                     '3.4382'],
            ['property_age *',         'INTEGER', 'Engineered: 2025 – yr_built',              '10'],
            ['renovation_flag *',      'INTEGER', 'Engineered: 1 if yr_renovated > 0',       '0'],
            ['log_sqft *',             'FLOAT',   'Engineered: ln(sqft_living + 1)',          '5.77'],
            ['bedroom_bathroom_ratio*','FLOAT',   'Engineered: bedrooms / bathrooms',        '1.33'],
        ]
    )
    body_no_indent(doc, '* Denotes features derived during the feature engineering step '
                        '(src/feature_engineering.py). These columns are not present in '
                        'the raw CSV but are created programmatically before model training.')

    doc.add_page_break()

    # ── Appendix D: Neighbourhood Configuration ─────────────────────
    section_heading(doc, 'Appendix D: Neighbourhood Configuration')
    body_no_indent(doc,
        'Table D.1 lists all twenty-six Nigerian neighbourhoods supported by the '
        'system, together with their geographic coordinates and price tier '
        'classifications. These values are defined in config.py and are used for '
        'both dataset generation and runtime inference (neighbourhood-to-coordinate '
        'lookup).')

    table_caption(doc, 'Table D.1: Neighbourhood Configuration (config.py – NEIGHBOURHOOD_COORDS)')
    add_table(doc,
        headers=['Neighbourhood', 'City', 'Price Tier', 'Base Price (₦/m²)', 'Latitude', 'Longitude'],
        rows=[
            ['Victoria Island', 'Lagos', 'Ultra Premium', '₦1,800,000', '6.4281', '3.4219'],
            ['Ikoyi', 'Lagos', 'Ultra Premium', '₦1,600,000', '6.4444', '3.4422'],
            ['Banana Island', 'Lagos', 'Ultra Premium', '₦2,500,000', '6.4527', '3.4361'],
            ['Lekki Phase 1', 'Lagos', 'Premium', '₦750,000', '6.4401', '3.5042'],
            ['Lekki Phase 2', 'Lagos', 'Premium', '₦600,000', '6.4573', '3.5387'],
            ['Ajah', 'Lagos', 'Mid-High', '₦300,000', '6.4682', '3.5796'],
            ['Gbagada', 'Lagos', 'Mid-High', '₦250,000', '6.5483', '3.3876'],
            ['Magodo', 'Lagos', 'Mid-High', '₦280,000', '6.5917', '3.3880'],
            ['Surulere', 'Lagos', 'Standard', '₦150,000', '6.5016', '3.3560'],
            ['Yaba', 'Lagos', 'Standard', '₦180,000', '6.5046', '3.3770'],
            ['Ikeja', 'Lagos', 'Standard', '₦160,000', '6.5956', '3.3382'],
            ['Ogba', 'Lagos', 'Standard', '₦130,000', '6.5836', '3.3447'],
            ['Maitama', 'Abuja', 'Ultra Premium', '₦1,200,000', '9.0723', '7.4892'],
            ['Asokoro', 'Abuja', 'Ultra Premium', '₦1,000,000', '9.0467', '7.5190'],
            ['Wuse 2', 'Abuja', 'Premium', '₦600,000', '9.0610', '7.4712'],
            ['Jabi', 'Abuja', 'Premium', '₦500,000', '9.0814', '7.4527'],
            ['Garki', 'Abuja', 'Mid-High', '₦350,000', '9.0465', '7.4836'],
            ['Gwarinpa', 'Abuja', 'Mid-High', '₦200,000', '9.1186', '7.4144'],
            ['Kubwa', 'Abuja', 'Standard', '₦120,000', '9.1395', '7.3491'],
            ['Lugbe', 'Abuja', 'Standard', '₦100,000', '9.0199', '7.4165'],
            ['GRA Phase 1', 'Port Harcourt', 'Premium', '₦500,000', '4.8117', '7.0199'],
            ['GRA Phase 2', 'Port Harcourt', 'Premium', '₦400,000', '4.8009', '7.0176'],
            ['Trans Amadi', 'Port Harcourt', 'Mid-High', '₦200,000', '4.8494', '7.0356'],
            ['Rumuola', 'Port Harcourt', 'Mid-High', '₦180,000', '4.8264', '7.0218'],
            ['Diobu', 'Port Harcourt', 'Standard', '₦100,000', '4.8228', '6.9986'],
            ['Eleme', 'Port Harcourt', 'Standard', '₦80,000', '4.7762', '7.0873'],
        ]
    )

    doc.add_page_break()

    # ── Appendix E: Glossary ────────────────────────────────────────
    section_heading(doc, 'Appendix E: Glossary of Technical Terms')

    terms = [
        ('Automated Valuation Model (AVM)',
         'A computer-based model that uses statistical or machine learning methods '
         'combined with property databases to estimate real estate values without '
         'requiring a human appraiser.'),
        ('Bootstrap Aggregating (Bagging)',
         'An ensemble learning technique that trains multiple models on different '
         'random subsets of the training data and averages their predictions '
         'to reduce variance. Forms the basis of the Random Forest algorithm.'),
        ('Coefficient of Determination (R²)',
         'A metric representing the proportion of variance in the target variable '
         'explained by the predictive model. Values range from 0 to 1, with '
         '1.0 indicating perfect predictions.'),
        ('ColumnTransformer',
         'A scikit-learn utility that applies different preprocessing transformations '
         '(e.g., scaling, encoding) to different columns of a dataset simultaneously, '
         'enabling a unified preprocessing pipeline.'),
        ('Cross-Validation (k-Fold)',
         'A model evaluation technique that partitions the dataset into k equal folds, '
         'trains on k-1 folds, and validates on the remaining fold, repeating the '
         'process k times to obtain a robust performance estimate.'),
        ('Flask',
         'A lightweight Python WSGI micro web framework used to build RESTful web '
         'applications and APIs.'),
        ('Gradient Boosting',
         'An ensemble learning technique that builds models sequentially, with each '
         'new model correcting the residual errors of its predecessor by minimising '
         'the gradient of a loss function.'),
        ('Hedonic Pricing Model',
         'An economic model that decomposes a heterogeneous product (such as a '
         'property) into its constituent characteristics and estimates the implicit '
         'price of each attribute.'),
        ('Joblib',
         'A Python library for efficient serialisation and deserialisation of '
         'large Python objects, commonly used to persist trained scikit-learn models.'),
        ('Mean Absolute Error (MAE)',
         'The average of the absolute differences between predicted and actual values. '
         'Measures average magnitude of prediction errors without directional bias.'),
        ('Mean Absolute Percentage Error (MAPE)',
         'The average of the absolute percentage differences between predicted and '
         'actual values. Provides a scale-independent measure of prediction accuracy.'),
        ('One-Hot Encoding',
         'A technique for converting categorical variables into a binary vector '
         'representation, where each category becomes a separate binary column.'),
        ('Overfitting',
         'A modelling error where a model learns noise in the training data too '
         'closely, resulting in poor generalisation to new, unseen data.'),
        ('Pipeline (scikit-learn)',
         'A scikit-learn object that chains multiple processing steps (e.g., '
         'preprocessing and model training) into a single object that can be '
         'fitted, evaluated, and serialised as a unit.'),
        ('Random Forest',
         'An ensemble machine learning algorithm that constructs multiple decision '
         'trees using bootstrap sampling and random feature selection, averaging '
         'their predictions to reduce variance.'),
        ('Root Mean Squared Error (RMSE)',
         'The square root of the average squared differences between predicted and '
         'actual values. Penalises larger errors more heavily than MAE.'),
        ('StandardScaler',
         'A scikit-learn preprocessing object that standardises numerical features '
         'to zero mean and unit variance.'),
        ('XGBoost (Extreme Gradient Boosting)',
         'A highly optimised, regularised gradient boosting implementation that '
         'supports parallel computation, native missing value handling, and '
         'L1/L2 regularisation to control overfitting.'),
    ]

    for term, definition in terms:
        para = doc.add_paragraph()
        _set_paragraph_format(para, space_before=4, space_after=4, line_spacing=1.5,
                               alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r1 = para.add_run(f'{term}: ')
        _set_run_font(r1, size=12, bold=True)
        r2 = para.add_run(definition)
        _set_run_font(r2, size=12)


# ─────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    set_document_margins(doc)

    # Set default document font in Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = FONT
    normal_style.font.size = Pt(12)
    nf = OxmlElement('w:rFonts')
    nf.set(qn('w:ascii'), FONT)
    nf.set(qn('w:hAnsi'), FONT)
    nf.set(qn('w:cs'), FONT)
    normal_style.element.get_or_add_rPr().insert(0, nf)

    print('Building Chapter Four...')
    build_chapter_four(doc)
    doc.add_page_break()

    print('Building Chapter Five...')
    build_chapter_five(doc)
    doc.add_page_break()

    print('Building Appendix...')
    build_appendix(doc)

    doc.save(OUTPUT_FILE)
    print(f'\nDocument saved: {OUTPUT_FILE}')
    size_kb = os.path.getsize(OUTPUT_FILE) // 1024
    print(f'File size: {size_kb} KB')


if __name__ == '__main__':
    main()
